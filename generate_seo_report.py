from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# Style setup
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for i, sz in enumerate([26, 20, 16, 13], 1):
    hs = doc.styles[f'Heading {i}']
    hs.font.size = Pt(sz)
    hs.font.color.rgb = RGBColor(0x2B, 0x2D, 0x42)
    hs.font.bold = True
    hs.font.name = 'Calibri'


def add_tbl(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tcPr = c._element.get_or_add_tcPr()
        sh = OxmlElement('w:shd')
        sh.set(qn('w:fill'), '2B2D42')
        sh.set(qn('w:val'), 'clear')
        tcPr.append(sh)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)


def add_bul(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)


# ═══════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Judicium Arbitration')
r.font.size = Pt(36)
r.font.bold = True
r.font.color.rgb = RGBColor(0x2B, 0x2D, 0x42)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('SEO & Technical Implementation Report')
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x2B, 0x2D, 0x42)

doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    'Comprehensive Documentation of Search Engine Optimization,\n'
    'Structured Data, Analytics Integration & Technical Architecture'
)
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

for _ in range(4):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Prepared for: Judicium Arbitration Stakeholders\nDate: March 2026')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc = [
    '1.  Executive Summary',
    '2.  Technical Architecture Overview',
    '3.  Structured Data & Schema Markup Implementation',
    '4.  Meta Tags & On-Page SEO Strategy',
    '5.  XML Sitemap Configuration',
    '6.  Robots.txt & Crawl Management',
    '7.  Page-by-Page SEO Breakdown',
    '8.  Build Process & Static Site Generation',
    '9.  Google Analytics (GA4) Integration',
    '10. Google Search Console Setup & Verification',
    '11. File Reference & Architecture Map',
    '12. Post-Deployment Validation Checklist',
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Cm(1)
    for run in p.runs:
        run.font.size = Pt(12)
doc.add_page_break()

# ═══════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'This document provides a comprehensive overview of the Search Engine Optimization (SEO) strategy '
    'and technical implementation carried out for the Judicium Arbitration website \u2014 a premier arbitration '
    'and alternative dispute resolution firm serving North India with over 20 years of experience and 500+ '
    'cases resolved. The goal of this initiative is to ensure that every page on the website is fully optimized '
    'for search engine discovery, indexing, and ranking, while simultaneously delivering an exceptional user '
    'experience across all devices.'
)

doc.add_paragraph(
    'The Judicium Arbitration website comprises 28 unique pages spanning practice areas, informational '
    'content, and legal service offerings across 8 cities in North India. Each page has been individually '
    'configured with targeted metadata, structured data markup, canonical URLs, and social media '
    'optimization tags to maximize visibility in Google Search results, Bing, and social media platforms.'
)

doc.add_heading('Key Achievements', level=2)

achievements = [
    ('28 Pages Fully Optimized: ',
     'Every page has unique title tags, meta descriptions, canonical URLs, and Open Graph tags tailored to its specific content and target keywords.'),
    ('4 Schema Types Implemented: ',
     'LegalService, Service, BreadcrumbList, and FAQPage structured data schemas are deployed across the platform to enable rich search results.'),
    ('20 Practice Area Pages: ',
     'Each practice area has dedicated SEO metadata, per-page JSON-LD structured data, and FAQ schema for rich result eligibility.'),
    ('Google Analytics (GA4) Integrated: ',
     'Full user behavior tracking is in place with Measurement ID G-LD76V5XMWC, capturing page views, session data, traffic sources, and user demographics.'),
    ('Google Search Console Ready: ',
     'The platform is configured for immediate Search Console verification, sitemap submission, and indexing monitoring.'),
    ('8-City Geographic SEO: ',
     'Structured data explicitly targets New Delhi, Gurgaon, Noida, Chandigarh, Jaipur, Panipat, Prayagraj, and Lucknow for local search visibility.'),
]
for b, t in achievements:
    add_bul(doc, t, b)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 2. TECHNICAL ARCHITECTURE
# ═══════════════════════════════════════════════════
doc.add_heading('2. Technical Architecture Overview', level=1)

doc.add_paragraph(
    'The Judicium Arbitration website is built using Next.js App Router, a modern React framework that '
    'provides first-class support for search engine optimization through its native Metadata API, file-based '
    'routing, and Static Site Generation capabilities.'
)

doc.add_heading('Why Next.js for SEO?', level=2)

for b, t in [
    ('Static Site Generation (SSG): ',
     'All 28 pages are pre-rendered at build time as static HTML files. Crawlers receive fully-rendered HTML immediately \u2014 no JavaScript execution required.'),
    ('Native Metadata API: ',
     'Built-in support for generating meta tags, Open Graph tags, Twitter cards through simple metadata object exports.'),
    ('File-Based Route Generation: ',
     'Critical SEO files (XML sitemap, robots.txt) are generated automatically from TypeScript source files.'),
]:
    add_bul(doc, t, b)

doc.add_heading('SEO Technology Stack', level=2)

add_tbl(doc,
        ['Layer', 'Source File', 'Output Generated'],
        [
            ['Structured Data', 'layout.tsx', 'LegalService JSON-LD on every page'],
            ['Structured Data', 'practice-areas/[slug]/page.tsx', 'Service + BreadcrumbList + FAQPage JSON-LD'],
            ['Meta Tags', 'layout.tsx', 'Global defaults (title template, OG, Twitter, robots)'],
            ['Meta Tags', 'Every page.tsx', 'Page-specific title, description, keywords, canonical, OG'],
            ['Sitemap', 'sitemap.ts', '/sitemap.xml with 28 URLs'],
            ['Robots', 'robots.ts', '/robots.txt with crawl directives'],
            ['Analytics', 'GoogleAnalytics.tsx', 'GA4 tracking on all pages'],
        ])

doc.add_heading('Metadata Inheritance Architecture', level=2)

doc.add_paragraph(
    'The SEO metadata follows a hierarchical inheritance model. The root layout file (layout.tsx) defines '
    'global defaults that apply to every page. Individual pages override only the properties they need to '
    'customize (title and description), while inheriting everything else (Open Graph images, Twitter card '
    'configuration, robots directives). When a child page exports a title like "Arbitration & Alternative '
    'Dispute Resolution," the template system automatically transforms it into "Arbitration & Alternative '
    'Dispute Resolution | Judicium Arbitration."'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 3. STRUCTURED DATA
# ═══════════════════════════════════════════════════
doc.add_heading('3. Structured Data & Schema Markup Implementation', level=1)

doc.add_heading('What is Structured Data?', level=2)
doc.add_paragraph(
    'Structured data is machine-readable code (in JSON-LD format) embedded in the HTML of web pages '
    'that helps search engines understand the content, context, and purpose of each page. It provides '
    'explicit, standardized information that search engines use to generate enhanced search results.'
)

doc.add_heading('Business Benefits of Structured Data', level=2)
for b, t in [
    ('Rich Snippets in Search Results: ',
     'FAQ structured data causes questions and answers to expand directly within Google search results, increasing click-through rates by an estimated 20\u201330%.'),
    ('Google Knowledge Panel: ',
     'The LegalService schema helps Google build a Knowledge Panel \u2014 displaying contact information, business hours, and description.'),
    ('Professional Breadcrumbs: ',
     'BreadcrumbList schema replaces raw URLs with clean navigation paths (Home > Practice Areas > Arbitration & ADR).'),
    ('Service Attribution: ',
     'Service schemas help Google accurately categorize each page, improving relevance for search queries.'),
    ('Local Pack Visibility: ',
     'Geographic data (8 cities, coordinates, address) enables visibility in Google Maps and the Local Pack.'),
]:
    add_bul(doc, t, b)

doc.add_heading('Schema 1: LegalService (Global)', level=2)
doc.add_paragraph(
    'This schema is embedded on every page and identifies the firm as a LegalService provider:'
)

add_tbl(doc,
        ['Field', 'Value', 'Purpose'],
        [
            ['@type', 'LegalService', 'Legal-specific rich results eligibility'],
            ['name', 'Judicium Arbitration', 'Knowledge Panel display'],
            ['telephone', '+91-9899686394', 'Click-to-call in mobile search results'],
            ['email', 'Judiciumarbitration@gmail.com', 'Contact entity information'],
            ['foundingDate', '2004', 'Trust signal \u2014 establishes longevity'],
            ['numberOfEmployees', 'minValue: 8', 'Business size context'],
            ['areaServed', '8 Cities across North India', 'Geographic targeting for local SEO'],
            ['address', 'New Delhi, Delhi, IN', 'Google Maps & Local Pack integration'],
            ['geo', 'Lat: 28.6139, Lng: 77.2090', 'Accurate map placement'],
            ['openingHours', 'Mon\u2013Fri, 09:00\u201318:00', 'Knowledge Panel & Google Maps display'],
            ['hasOfferCatalog', '6 core services', 'Service-specific query matching'],
        ])

doc.add_heading('Schema 2: Service (Per Practice Area)', level=2)
doc.add_paragraph(
    'Each of the 20 practice area pages includes a Service schema specifying the service name, description, '
    'provider (Judicium Arbitration as a LegalService entity), area served (8 cities), and canonical URL.'
)

doc.add_heading('Schema 3: BreadcrumbList (All Detail Pages)', level=2)
doc.add_paragraph(
    'Every practice area page includes a BreadcrumbList schema defining the navigation hierarchy '
    '(Home > Practice Areas > Current Area). Google shows clean breadcrumb trails instead of raw URLs.'
)

doc.add_heading('Schema 4: FAQPage (Conditional)', level=2)
doc.add_paragraph(
    'Pages with FAQ content are annotated with the FAQPage schema. FAQ rich results can occupy '
    '3\u20134x more SERP space than standard listings, pushing competitors below the fold.'
)

doc.add_heading('Schema Coverage Summary', level=2)
add_tbl(doc,
        ['Schema Type', 'Where Applied', 'Pages', 'Key Benefit'],
        [
            ['LegalService', 'Every page (via layout)', '28', 'Knowledge Panel, brand identity, local SEO'],
            ['Service', 'Practice area pages', '20', 'Service-specific rich results'],
            ['BreadcrumbList', 'Practice area detail pages', '20', 'Professional breadcrumb display'],
            ['FAQPage', 'Pages with FAQ content', 'Varies', 'Expandable FAQ rich results'],
        ])

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 4. META TAGS
# ═══════════════════════════════════════════════════
doc.add_heading('4. Meta Tags & On-Page SEO Strategy', level=1)

doc.add_paragraph(
    'Meta tags are HTML elements that provide structured information about a web page to search engines '
    'and social media platforms. They control what Google displays in search results, how pages appear '
    'when shared on social media, how search engines crawl the page, and how duplicate content is handled.'
)

doc.add_heading('4.1 Title Tags', level=2)
doc.add_paragraph(
    'Every page has a unique, keyword-optimized title tag. The root layout defines a template '
    '"%s | Judicium Arbitration" that automatically appends the brand name. The homepage uses '
    '"Judicium Arbitration | Leading Arbitration Services in North India" as the default.'
)

doc.add_heading('4.2 Meta Descriptions', level=2)
doc.add_paragraph(
    'Each page has a unique meta description optimized to 155\u2013160 characters. The global description '
    'highlights premier arbitration services across 8 cities with 20+ years experience and 500+ cases resolved. '
    'Individual practice area pages override this with content-specific descriptions.'
)

doc.add_heading('4.3 Enhanced Keywords (18 terms)', level=2)
doc.add_paragraph(
    'arbitration, ADR, alternative dispute resolution, arbitration India, arbitration lawyers Delhi, '
    'commercial arbitration, international arbitration, dispute resolution, legal services Delhi, '
    'law firm North India, Delhi, Gurgaon, Noida, Chandigarh, Jaipur, Panipat, Prayagraj, Lucknow'
)

doc.add_heading('4.4 Canonical URLs', level=2)
doc.add_paragraph(
    'Every page specifies a canonical URL to prevent duplicate content issues from URL variations '
    '(trailing slashes, capitalizations, tracking parameters).'
)

doc.add_heading('4.5 Robots Directives', level=2)
add_tbl(doc,
        ['Directive', 'Value', 'Effect'],
        [
            ['index', 'true', 'Pages appear in search results'],
            ['follow', 'true', 'Crawlers follow links to discover more content'],
            ['max-image-preview', 'large', 'Large image thumbnails in search results & Discover feed'],
            ['max-snippet', '-1 (unlimited)', 'Extended text snippets for more SERP real estate'],
            ['max-video-preview', '-1 (unlimited)', 'Full video previews (future-proofing)'],
        ])

doc.add_heading('4.6 Open Graph Tags', level=2)
doc.add_paragraph(
    'Open Graph tags control social sharing appearance on Facebook, LinkedIn, WhatsApp. Every page '
    'includes site name, title, description, preview image (logo.jpeg 800x600), canonical URL, content type, '
    'and locale (en_IN). These ensure professional, branded previews that encourage engagement.'
)

doc.add_heading('4.7 Twitter Card Tags', level=2)
doc.add_paragraph(
    'Configured as summary_large_image \u2014 the large card format showing a big image above text for '
    'maximum visual prominence on Twitter/X.'
)

doc.add_heading('4.8 Authors & E-E-A-T Signals', level=2)
doc.add_paragraph(
    'Author, creator, and publisher meta tags set to "Judicium Arbitration" for E-E-A-T (Experience, '
    "Expertise, Authoritativeness, Trustworthiness) \u2014 Google's quality framework that heavily impacts "
    'legal website rankings.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 5. SITEMAP
# ═══════════════════════════════════════════════════
doc.add_heading('5. XML Sitemap Configuration', level=1)

doc.add_paragraph(
    'The XML sitemap lists every URL on the website with metadata about last modification, change frequency, '
    "and relative importance. Served at https://judicium-arbitration.com/sitemap.xml, it is auto-generated from the platform's data sources."
)

doc.add_heading('Why a Sitemap is Essential', level=2)
for b, t in [
    ('Comprehensive Discoverability: ', 'Ensures search engines discover every page, including 20 dynamically generated practice area pages.'),
    ('Crawl Budget Optimization: ', 'Priority and frequency metadata helps search engines focus on the most important content first.'),
    ('Faster Indexing: ', 'New practice areas appear in the sitemap immediately after the next build.'),
]:
    add_bul(doc, t, b)

doc.add_heading('Priority Strategy', level=2)
add_tbl(doc,
        ['Priority', 'Pages', 'Rationale'],
        [
            ['1.0', 'Homepage', 'Most authoritative page; all link equity flows from here'],
            ['0.9', 'Practice Areas listing', 'Primary conversion page targeting high-volume queries'],
            ['0.8', 'About, Contact', 'Trust-building pages important for conversions'],
            ['0.7', '20 Practice area detail pages', 'Valuable service pages, secondary to listing'],
            ['0.3', 'Privacy Policy, Terms, Disclaimer', 'Required legal pages, not ranking targets'],
        ])

doc.add_heading('Change Frequency Strategy', level=2)
add_tbl(doc,
        ['Frequency', 'Applied To', 'Rationale'],
        [
            ['Weekly', 'Homepage, Practice Areas listing', 'May be updated frequently with new offerings'],
            ['Monthly', 'About, Contact, practice area pages', 'Relatively stable with periodic updates'],
            ['Yearly', 'Privacy Policy, Terms, Disclaimer', 'Legal boilerplate that rarely changes'],
        ])

doc.add_heading('Complete URL Inventory: 28 Pages', level=2)
add_tbl(doc,
        ['Category', 'Number of URLs'],
        [
            ['Static pages (Home, About, Practice Areas, Contact)', '4'],
            ['Individual Practice Area detail pages', '20'],
            ['Legal pages (Privacy, Terms, Disclaimer)', '3'],
            ['Total', '28 (auto-expands as practice areas are added)'],
        ])

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 6. ROBOTS.TXT
# ═══════════════════════════════════════════════════
doc.add_heading('6. Robots.txt & Crawl Management', level=1)

doc.add_paragraph(
    'The robots.txt file served at https://judicium-arbitration.com/robots.txt communicates crawling '
    'instructions to search engine bots.'
)

doc.add_heading('Configuration Details', level=2)
add_tbl(doc,
        ['Rule', 'Value', 'Purpose'],
        [
            ['User-Agent', '* (all bots)', 'Universal crawling rules for all search engines'],
            ['Allow', '/', 'Permits crawling of all public content'],
            ['Disallow', '/api/', 'Blocks API routes returning JSON, not HTML'],
            ['Sitemap', 'https://judicium-arbitration.com/sitemap.xml', 'Directs crawlers to sitemap for URL discovery'],
        ])

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 7. PAGE-BY-PAGE
# ═══════════════════════════════════════════════════
doc.add_heading('7. Page-by-Page SEO Breakdown', level=1)

doc.add_heading('7.1 Static Pages', level=2)
doc.add_paragraph('The following static pages each have individually crafted metadata, canonical URLs, and inherit the global LegalService structured data schema:')

add_tbl(doc,
        ['Page', 'URL Path', 'Title'],
        [
            ['Home', '/', 'Judicium Arbitration | Leading Arbitration Services in North India'],
            ['About', '/about', 'About Us | Judicium Arbitration - Leading Legal Experts in North India'],
            ['Practice Areas', '/practice-areas', 'Practice Areas | Judicium Arbitration - Expert Legal Services'],
            ['Contact', '/contact', 'Contact Us | Judicium Arbitration - Get Legal Consultation'],
            ['Privacy Policy', '/privacy-policy', 'Privacy Policy | Judicium Arbitration'],
            ['Terms of Service', '/terms-of-service', 'Terms of Service | Judicium Arbitration'],
            ['Disclaimer', '/disclaimer', 'Disclaimer | Judicium Arbitration'],
        ])

doc.add_heading('7.2 Dynamic Practice Area Pages (20 pages)', level=2)
doc.add_paragraph('Each page receives unique metadata and Service + BreadcrumbList + FAQPage JSON-LD schemas:')

practice_areas = [
    'Arbitration & Alternative Dispute Resolution', 'Banking & Finance', 'Capital Markets',
    'Competition and Anti-Trust', 'Corporate Mergers & Acquisitions', 'Corporate and Commercial',
    'Data Privacy & Cybersecurity', 'Real Estate & Urban Development', 'Intellectual Property Rights',
    'Insolvency & Bankruptcy', 'Labour and Employment', 'Dispute Resolution \u2013 Litigation',
    'Technology, Media & Telecommunication', 'Projects & Infrastructure',
    'Private Equity and Investment Funds', 'White Collar Crime', 'Healthcare & Lifesciences',
    'International Trade/WTO', 'Insurance', 'Defence, Aviation & Space',
]
for pa in practice_areas:
    add_bul(doc, pa)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 8. BUILD PROCESS
# ═══════════════════════════════════════════════════
doc.add_heading('8. Build Process & Static Site Generation', level=1)

doc.add_paragraph('When the website is built for deployment, the following process occurs:')

steps = [
    'Next.js reads all page files and identifies every route.',
    'For practice areas, generateStaticParams() retrieves all URL slugs from practiceAreas.ts.',
    'Each page is pre-rendered as complete static HTML with meta tags in <head>.',
    'sitemap.ts generates /sitemap.xml with all 28 URLs.',
    'robots.ts generates /robots.txt with crawl directives.',
    'JSON-LD structured data is injected as <script> tags into every relevant page.',
    'Google Analytics GA4 component is included via root layout on every page.',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('Why SSG Matters for SEO', level=2)
for b, t in [
    ('Instant Crawlability: ', 'Crawlers receive fully-rendered HTML immediately \u2014 no JavaScript dependency.'),
    ('Maximum Page Speed: ', 'Static pages served from CDN achieve fastest possible load times, boosting Core Web Vitals.'),
    ('Reliable Content Delivery: ', 'No risk of rendering failures or API timeouts preventing crawlers from seeing content.'),
    ('Optimal Core Web Vitals: ', 'LCP, FID, and CLS scores maximized with minimal client-side processing.'),
]:
    add_bul(doc, t, b)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 9. GA4
# ═══════════════════════════════════════════════════
doc.add_heading('9. Google Analytics (GA4) Integration', level=1)

doc.add_paragraph(
    'Google Analytics 4 has been integrated to provide comprehensive insights into user behavior, '
    'traffic patterns, and engagement metrics across all pages.'
)

doc.add_heading('Implementation Details', level=2)
add_tbl(doc,
        ['Configuration', 'Value'],
        [
            ['GA4 Measurement ID', 'G-LD76V5XMWC'],
            ['Analytics Dashboard', 'analytics.google.com'],
            ['Loading Strategy', 'afterInteractive (does not block page rendering)'],
            ['Environment Variable', 'NEXT_PUBLIC_GA_MEASUREMENT_ID'],
            ['Scope', 'All pages (loaded via root layout)'],
        ])

doc.add_heading('Technical Decisions', level=2)
for b, t in [
    ('afterInteractive Loading: ', 'Script loads after page is interactive, never blocking rendering or hurting Core Web Vitals.'),
    ('Environment-Based Config: ', 'Measurement ID stored as env variable, allowing different properties per environment.'),
    ('Graceful Degradation: ', 'If env variable is not set, the component renders nothing with no errors.'),
]:
    add_bul(doc, t, b)

doc.add_heading('Default Tracking', level=2)
for t in [
    'Page views across all 28 pages',
    'Scroll depth per page',
    'Outbound link clicks',
    'Session duration and engagement',
    'Traffic source attribution (Google, social, direct, referral)',
    'User demographics (country, city, device, browser, OS)',
]:
    add_bul(doc, t)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 10. SEARCH CONSOLE
# ═══════════════════════════════════════════════════
doc.add_heading('10. Google Search Console Setup & Verification', level=1)

doc.add_paragraph(
    'Google Search Console provides direct insight into how Google perceives, crawls, and indexes the website.'
)

doc.add_heading('Setup Procedure', level=2)
for i, s in enumerate([
    'Navigate to search.google.com/search-console.',
    'Add property https://judicium-arbitration.com.',
    'Select HTML tag verification method.',
    'Copy the verification code from Google.',
    'Set NEXT_PUBLIC_GOOGLE_SITE_VERIFICATION in .env.local.',
    'Deploy the updated configuration.',
    'Return to Search Console and click Verify.',
    'Submit sitemap.xml in the Sitemaps section.',
], 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('Post-Verification Actions', level=2)
add_tbl(doc,
        ['Action', 'Location in GSC', 'Purpose'],
        [
            ['Submit Sitemap', 'Sitemaps section', 'Discover and process all 28 URLs'],
            ['Monitor Index Coverage', 'Pages > Indexing', 'Identify crawl errors'],
            ['Validate Structured Data', 'Enhancements tab', 'Confirm schemas are error-free'],
            ['Track Performance', 'Performance tab', 'Monitor impressions, clicks, CTR, position'],
            ['Core Web Vitals', 'Experience > Core Web Vitals', "Ensure LCP, FID, CLS meet thresholds"],
            ['Mobile Usability', 'Experience > Mobile Usability', "Confirm mobile-friendly assessment"],
        ])

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 11. FILE REFERENCE
# ═══════════════════════════════════════════════════
doc.add_heading('11. File Reference & Architecture Map', level=1)

doc.add_heading('11.1 SEO & Analytics Core Files', level=2)
add_tbl(doc,
        ['File', 'Path', 'Purpose'],
        [
            ['Root Layout', 'src/app/layout.tsx', 'Global metadata, JSON-LD, viewport, GA4, Search Console tag'],
            ['Google Analytics', 'src/components/GoogleAnalytics.tsx', 'GA4 tracking script (client-side)'],
            ['Sitemap', 'src/app/sitemap.ts', 'Generates /sitemap.xml with 28 URLs'],
            ['Robots', 'src/app/robots.ts', 'Generates /robots.txt'],
            ['Env Variables', '.env.local', 'GA4 ID, Search Console verification code'],
        ])

doc.add_heading('11.2 Pages with Custom SEO Metadata', level=2)
add_tbl(doc,
        ['Page', 'File Path'],
        [
            ['Home', 'src/app/page.tsx'],
            ['About', 'src/app/about/page.tsx'],
            ['Practice Areas (listing)', 'src/app/practice-areas/page.tsx'],
            ['Practice Area (detail)', 'src/app/practice-areas/[slug]/page.tsx'],
            ['Contact', 'src/app/contact/page.tsx'],
            ['Privacy Policy', 'src/app/privacy-policy/page.tsx'],
            ['Terms of Service', 'src/app/terms-of-service/page.tsx'],
            ['Disclaimer', 'src/app/disclaimer/page.tsx'],
        ])

doc.add_heading('11.3 Data File', level=2)
add_tbl(doc,
        ['Data File', 'Path', 'Drives'],
        [
            ['Practice Areas', 'src/data/practiceAreas.ts', 'Practice area pages, sitemap, FAQ schemas'],
        ])

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 12. VALIDATION CHECKLIST
# ═══════════════════════════════════════════════════
doc.add_heading('12. Post-Deployment Validation Checklist', level=1)

doc.add_paragraph('After deploying to production, complete the following validation steps:')

checks = [
    ('1. Google Rich Results Test',
     'Test any page URL to verify structured data (FAQ, Breadcrumbs, LegalService) is valid and eligible for rich results.'),
    ('2. Schema.org Validator',
     'Validate JSON-LD syntax of all structured data schemas against Schema.org specification.'),
    ('3. Facebook Sharing Debugger',
     'Test how pages appear when shared on Facebook and Open Graph-compatible platforms.'),
    ('4. Twitter Card Validator',
     'Verify Twitter Card tags produce the expected large image summary card.'),
    ('5. Google PageSpeed Insights',
     'Analyze Core Web Vitals (LCP, FID, CLS) and performance scores for mobile and desktop.'),
    ('6. Google Search Console',
     'Submit sitemap, verify ownership, monitor index coverage and search performance.'),
    ('7. Robots.txt Verification',
     'Confirm /robots.txt returns expected crawl directives and sitemap reference.'),
    ('8. Sitemap Verification',
     'Confirm /sitemap.xml lists all 28 URLs with correct priorities and frequencies.'),
    ('9. GA4 Verification',
     'Visit site and check Realtime report in GA4 to confirm tracking with G-LD76V5XMWC.'),
]
for title, desc in checks:
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(11)
    p2 = doc.add_paragraph(desc)
    p2.paragraph_format.left_indent = Cm(1)

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "This documentation covers the complete SEO implementation including Google Analytics (GA4) "
    "and Google Search Console integration as of March 2026. All changes follow Google\u2019s Webmaster "
    "Guidelines and Schema.org specifications."
)
r.font.size = Pt(10)
r.italic = True
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ═══ SAVE ═══
out = r'C:\Users\bolis\Desktop\LAWWEBSITE\judicium_arbitration\Judicium_Arbitration_SEO_Report.docx'
doc.save(out)
print(f'SAVED: {out}')
print(f'Size: {os.path.getsize(out)} bytes')
